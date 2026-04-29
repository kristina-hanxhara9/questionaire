"""Generate the bundled sample assets (Excel guide, PDF reference, DOCX template).

Run this once after installing dependencies on Windows:

    py -m pip install -e .
    py scripts\\generate_samples.py

It writes four files into ``src/mcp_market_research/samples/``:

- ``sample_channel_guide.xlsx`` — legacy guidance workbook (3 channels:
  social_media, email, in_store). The ``Questionnaire Generator``
  (legacy single-agent) flow uses this.
- ``sample_modules.xlsx``       — module workbook (core + industry__opticians +
  unique__opticians__<channel> sheets) for the orchestrator + composer +
  reviewer + translator multi-agent flow.
- ``sample_template.pdf``       — visual layout reference with placeholders.
- ``sample_template.docx``      — docxtpl Jinja2 template that mirrors the PDF.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)

SAMPLES = Path(__file__).resolve().parents[1] / "src" / "mcp_market_research" / "samples"


def build_excel(path: Path) -> None:
    wb = Workbook()
    wb.remove(wb.active)

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="1F4E78")
    wrap = Alignment(wrap_text=True, vertical="top")
    headers = ["section", "question_type", "guidance", "example", "required", "notes"]

    channels: dict[str, list[list[object]]] = {
        "social_media": [
            ["Awareness", "single_choice",
             "Ask which platforms the respondent uses most for discovering brands like the client.",
             "Which platforms do you use most often to discover new brands?",
             "yes", ""],
            ["Awareness", "open_text",
             "Probe how the respondent first heard about the client's category.",
             "How did you first hear about <category>?",
             "no", ""],
            ["Engagement", "rating",
             "Measure 1-5 how engaging the respondent finds the client's social content.",
             "Rate the client's social content (1=boring, 5=very engaging).",
             "yes", "Use a 5-point scale, anchor labels."],
            ["Engagement", "multi_choice",
             "Which content formats drive the most engagement.",
             "Which formats do you engage with? (select all)",
             "yes", "Options: short video, long video, image, carousel, story, livestream."],
            ["Conversion", "yes_no",
             "Has the respondent ever purchased after seeing a social ad from the client.",
             "Have you ever bought from <client> after seeing one of their social posts/ads?",
             "yes", ""],
        ],
        "email": [
            ["Subscription", "single_choice",
             "How the respondent joined the email list.",
             "How did you originally sign up for our emails?",
             "yes", ""],
            ["Content", "rating",
             "Perceived usefulness of email content on a 1-5 scale.",
             "How useful do you find the emails? (1=not at all, 5=extremely)",
             "yes", ""],
            ["Frequency", "single_choice",
             "Preferred email cadence.",
             "What is your preferred email frequency?",
             "yes", "Options: daily, weekly, biweekly, monthly, less."],
            ["Content", "multi_choice",
             "Topics the respondent wants to see more of.",
             "Which topics would you like to see more of?",
             "no", ""],
            ["Loyalty", "open_text",
             "What would make the respondent recommend the newsletter.",
             "What single change would make you recommend our newsletter?",
             "no", ""],
        ],
        "in_store": [
            ["Visit", "single_choice",
             "Frequency of visits to the physical location.",
             "How often do you visit our store?",
             "yes", ""],
            ["Experience", "rating",
             "Overall in-store experience rating, 1-5.",
             "Rate your overall experience today (1=poor, 5=excellent).",
             "yes", ""],
            ["Staff", "rating",
             "Helpfulness of staff during the visit.",
             "How helpful were our staff today?",
             "yes", ""],
            ["Product", "open_text",
             "Anything the respondent wishes had been in stock.",
             "Was there anything you were looking for that we didn't have?",
             "no", ""],
            ["Loyalty", "yes_no",
             "Whether the respondent intends to return.",
             "Do you plan to visit us again in the next 3 months?",
             "yes", ""],
        ],
    }

    for sheet_name, rows in channels.items():
        ws = wb.create_sheet(sheet_name)
        ws.append(headers)
        for cell in ws[1]:
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = wrap
        for row in rows:
            ws.append(row)
        for col_letter, width in zip("ABCDEF", (16, 16, 50, 50, 10, 30), strict=False):
            ws.column_dimensions[col_letter].width = width
        for row_cells in ws.iter_rows(min_row=2):
            for cell in row_cells:
                cell.alignment = wrap

    notes = wb.create_sheet("_meta")
    notes["A1"] = "This workbook is the bundled sample. Add or modify channels by adding sheets."
    notes["A2"] = "Sheets prefixed with '_' (like this one) are skipped by the parser."

    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(path)


def build_pdf(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title="Market Research Questionnaire Template",
    )

    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], fontSize=18, spaceAfter=12)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], fontSize=14, spaceAfter=8)
    body = ParagraphStyle("body", parent=styles["BodyText"], fontSize=11, leading=14)

    story: list[object] = [
        Paragraph("{{COMPANY_NAME}}", h1),
        Paragraph("{{QUESTIONNAIRE_TITLE}}", h2),
        Paragraph(
            "Language: {{LANGUAGE}} &nbsp;&nbsp; Country: {{COUNTRY}} "
            "&nbsp;&nbsp; Generated: {{GENERATED_DATE}}",
            body,
        ),
        Spacer(1, 0.6 * cm),
        Paragraph("Introduction", h2),
        Paragraph(
            "Thank you for taking part in this short survey. Your answers help us "
            "improve our products and services. Responses are confidential and used "
            "only in aggregate.",
            body,
        ),
        Spacer(1, 0.4 * cm),
        Paragraph("&lt;&lt;SECTIONS_START&gt;&gt;", body),
        Paragraph("Section heading", h2),
        Paragraph("1. Sample question text here. *", body),
        Paragraph("[ ] Option A &nbsp;&nbsp; [ ] Option B &nbsp;&nbsp; [ ] Option C", body),
        Paragraph("2. Another question here.", body),
        Paragraph("&lt;&lt;SECTIONS_END&gt;&gt;", body),
        Spacer(1, 0.6 * cm),
        Paragraph("Closing", h2),
        Paragraph(
            "Thank you for your time. If you have additional comments, please share "
            "them with us at {{COMPANY_NAME}} support.",
            body,
        ),
    ]
    doc.build(story)


def build_docx(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("{{ company_name }}")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run("{{ questionnaire_title }}")
    sub_run.bold = True
    sub_run.font.size = Pt(14)

    meta = doc.add_paragraph()
    meta.add_run(
        "Language: {{ language }}    Country: {{ country_code or '-' }}    "
        "Generated: {{ generated_date }}"
    ).italic = True

    doc.add_paragraph(
        "Thank you for taking part in this short survey. Your answers help us "
        "improve our products and services. Responses are confidential and used "
        "only in aggregate."
    )

    doc.add_paragraph("{% for section in sections %}")
    heading = doc.add_paragraph()
    h_run = heading.add_run("{{ section.heading }}")
    h_run.bold = True
    h_run.font.size = Pt(13)

    doc.add_paragraph("{% for q in section.questions %}")
    doc.add_paragraph(
        "{{ loop.index }}. {{ q.text }}{% if q.required %} *{% endif %}"
    )
    doc.add_paragraph(
        "{% if q.options %}{% for opt in q.options %}[ ] {{ opt }}    "
        "{% endfor %}{% endif %}"
    )
    doc.add_paragraph("{% endfor %}")
    doc.add_paragraph("{% endfor %}")

    doc.add_paragraph(
        "Thank you for your time. If you have additional comments, please share "
        "them with us at {{ company_name }} support."
    )

    doc.save(str(path))


def build_modules_workbook(path: Path) -> None:
    """Build a sample module workbook for the orchestrator workflow.

    Sheets:
      core                              — channel-agnostic, industry-agnostic (~12 demo rows)
      industry__opticians               — opticians-specific, channel-agnostic (~10 demo rows)
      unique__opticians__social_media   — opticians × social media (~6 rows)
      unique__opticians__email          — opticians × email (~6 rows)
      unique__opticians__in_store       — opticians × in-store (~6 rows)

    Demo content uses smaller counts than the production target (84 / 74 / 6) so
    the file is readable. The parser doesn't care about counts; only structure.
    """
    wb = Workbook()
    wb.remove(wb.active)

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="1F4E78")
    wrap = Alignment(wrap_text=True, vertical="top")
    headers = ["id", "section", "question_text", "question_type", "options", "required", "tags", "notes"]

    sheets: dict[str, list[list[object]]] = {
        "core": [
            ["core_001", "Awareness", "How did you first hear about <COMPANY>?", "single_choice",
             "Search engine|Social media|Word of mouth|Advertising|Other", "no", "awareness,top_of_funnel", ""],
            ["core_002", "Awareness", "Which of the following brands have you heard of?", "multi_choice",
             "<COMPANY>|Competitor A|Competitor B|None of the above", "no", "awareness,brand", ""],
            ["core_003", "Perception", "How would you rate <COMPANY> overall?", "rating",
             "1|2|3|4|5", "no", "perception,brand,quality", "5-point scale, 1=poor, 5=excellent"],
            ["core_004", "Perception", "What three words best describe <COMPANY>?", "open_text",
             "", "no", "perception,brand,open", ""],
            ["core_005", "Loyalty", "How likely are you to recommend <COMPANY> to a friend?", "rating",
             "0|1|2|3|4|5|6|7|8|9|10", "no", "nps,referral,loyalty", "Standard NPS 0-10"],
            ["core_006", "Loyalty", "How likely are you to use <COMPANY> again in the next 12 months?", "rating",
             "1|2|3|4|5", "no", "retention,loyalty", ""],
            ["core_007", "Purchase", "Which factors most influence your decision when buying in this category?", "multi_choice",
             "Price|Quality|Brand reputation|Convenience|Recommendations|Other", "no", "purchase,decision", ""],
            ["core_008", "Purchase", "How sensitive are you to price changes from <COMPANY>?", "rating",
             "1|2|3|4|5", "no", "purchase,pricing", ""],
            ["core_009", "Feedback", "What single change would most improve your experience?", "open_text",
             "", "no", "feedback,open", ""],
            ["core_010", "Demographics", "What is your age group?", "single_choice",
             "Under 18|18-24|25-34|35-44|45-54|55-64|65+", "no", "demographic", ""],
            ["core_011", "Demographics", "Which gender do you most identify with?", "single_choice",
             "Female|Male|Non-binary|Prefer not to say|Self-describe", "no", "demographic", ""],
            ["core_012", "Consent", "May we contact you for follow-up research?", "yes_no",
             "", "yes", "consent,pii", "Required by GDPR/CCPA for follow-up contact"],
        ],
        "industry__opticians": [
            ["opt_std_001", "Vision needs", "When did you last have an eye exam?", "single_choice",
             "Within the last 6 months|6-12 months ago|1-2 years ago|2+ years ago|Never", "no", "behavior,frequency", ""],
            ["opt_std_002", "Vision needs", "Do you currently wear prescription eyewear?", "single_choice",
             "Glasses only|Contacts only|Both|Neither", "no", "screener,behavior", ""],
            ["opt_std_003", "Purchase", "Where did you last purchase eyewear?", "single_choice",
             "<COMPANY>|Competitor A|Competitor B|Online retailer|Department store|Other", "no", "purchase,channel", ""],
            ["opt_std_004", "Purchase", "How important is in-store eye exam availability when choosing an optician?", "rating",
             "1|2|3|4|5", "no", "purchase,decision,perception", ""],
            ["opt_std_005", "Perception", "How would you rate the selection of frames at <COMPANY>?", "rating",
             "1|2|3|4|5", "no", "perception,attribute,quality", ""],
            ["opt_std_006", "Perception", "How would you rate the lens quality at <COMPANY>?", "rating",
             "1|2|3|4|5", "no", "perception,attribute,quality", ""],
            ["opt_std_007", "Service", "How knowledgeable were the staff who helped you?", "rating",
             "1|2|3|4|5", "no", "perception,service", ""],
            ["opt_std_008", "Service", "How likely are you to use <COMPANY>'s eye exam service again?", "rating",
             "1|2|3|4|5", "no", "retention,service", ""],
            ["opt_std_009", "Loyalty", "Would you recommend our optician practice to family or friends?", "rating",
             "0|1|2|3|4|5|6|7|8|9|10", "no", "nps,referral,loyalty", "Industry-specific NPS variant"],
            ["opt_std_010", "Pricing", "How fair did you find the pricing for your last purchase?", "rating",
             "1|2|3|4|5", "no", "purchase,pricing", ""],
        ],
        "unique__opticians__social_media": [
            ["opt_sm_001", "Channel awareness", "On which social platforms have you seen <COMPANY> content?", "multi_choice",
             "Instagram|Facebook|TikTok|YouTube|Pinterest|None", "no", "awareness,channel,behavior", ""],
            ["opt_sm_002", "Engagement", "How engaging do you find <COMPANY>'s eyewear styling content?", "rating",
             "1|2|3|4|5", "no", "behavior,channel,perception", ""],
            ["opt_sm_003", "Engagement", "Which content formats from <COMPANY> do you engage with most?", "multi_choice",
             "Try-on videos|Before/after photos|Style guides|Eye health tips|Promotions|None", "no", "behavior,channel", ""],
            ["opt_sm_004", "Conversion", "Have you ever booked an eye exam after seeing <COMPANY>'s social content?", "yes_no",
             "", "no", "purchase,conversion,channel", ""],
            ["opt_sm_005", "Conversion", "Have you purchased eyewear after seeing <COMPANY>'s social content?", "yes_no",
             "", "no", "purchase,conversion,channel", ""],
            ["opt_sm_006", "Voice", "How would you describe <COMPANY>'s social media tone?", "open_text",
             "", "no", "perception,brand,channel,open", ""],
        ],
        "unique__opticians__email": [
            ["opt_em_001", "Subscription", "How did you originally sign up for <COMPANY>'s emails?", "single_choice",
             "After eye exam|At checkout|On the website|From a promotion|Don't remember", "no", "behavior,channel", ""],
            ["opt_em_002", "Content", "How useful do you find <COMPANY>'s eye-care emails?", "rating",
             "1|2|3|4|5", "no", "perception,channel", ""],
            ["opt_em_003", "Content", "Which topics from <COMPANY> would you like to see more of?", "multi_choice",
             "Frame trends|Lens technology|Eye health tips|Promotions|Product launches|None", "no", "behavior,channel", ""],
            ["opt_em_004", "Frequency", "What is your preferred email frequency from <COMPANY>?", "single_choice",
             "Weekly|Biweekly|Monthly|Quarterly|Less", "no", "behavior,channel", ""],
            ["opt_em_005", "Conversion", "Have you booked an eye exam from a <COMPANY> email in the last year?", "yes_no",
             "", "no", "purchase,conversion,channel", ""],
            ["opt_em_006", "Loyalty", "What single change would make you recommend <COMPANY>'s newsletter?", "open_text",
             "", "no", "feedback,channel,open", ""],
        ],
        "unique__opticians__in_store": [
            ["opt_is_001", "Visit", "How often do you visit a <COMPANY> location?", "single_choice",
             "Weekly|Monthly|Every few months|Once a year|Less than yearly", "no", "behavior,frequency", ""],
            ["opt_is_002", "Experience", "How welcoming did you find the <COMPANY> store on your last visit?", "rating",
             "1|2|3|4|5", "no", "perception,service,channel", ""],
            ["opt_is_003", "Try-on", "How helpful was the frame try-on process?", "rating",
             "1|2|3|4|5", "no", "behavior,service,channel", ""],
            ["opt_is_004", "Wait", "How would you rate the wait time for your eye exam?", "rating",
             "1|2|3|4|5", "no", "perception,service,channel", ""],
            ["opt_is_005", "Inventory", "Was the frame style or lens option you wanted in stock?", "yes_no",
             "", "no", "behavior,inventory,channel", ""],
            ["opt_is_006", "Loyalty", "How likely are you to return to this <COMPANY> location?", "rating",
             "1|2|3|4|5", "no", "retention,loyalty,channel", ""],
        ],
    }

    for sheet_name, rows in sheets.items():
        ws = wb.create_sheet(sheet_name)
        ws.append(headers)
        for cell in ws[1]:
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = wrap
        for row in rows:
            ws.append(row)
        for col_letter, width in zip("ABCDEFGH", (14, 18, 60, 16, 40, 10, 30, 24), strict=False):
            ws.column_dimensions[col_letter].width = width
        for row_cells in ws.iter_rows(min_row=2):
            for cell in row_cells:
                cell.alignment = wrap

    notes = wb.create_sheet("_meta")
    notes["A1"] = "Module workbook: core + industry__<industry> + unique__<industry>__<channel>."
    notes["A2"] = "Sheets prefixed with '_' (like this one) are skipped by the parser."
    notes["A3"] = "Add new industries/channels by adding sheets that follow the naming convention."

    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(path)


def main() -> None:
    SAMPLES.mkdir(parents=True, exist_ok=True)
    excel_path = SAMPLES / "sample_channel_guide.xlsx"
    modules_path = SAMPLES / "sample_modules.xlsx"
    pdf_path = SAMPLES / "sample_template.pdf"
    docx_path = SAMPLES / "sample_template.docx"

    build_excel(excel_path)
    print(f"wrote {excel_path}")
    build_modules_workbook(modules_path)
    print(f"wrote {modules_path}")
    build_pdf(pdf_path)
    print(f"wrote {pdf_path}")
    build_docx(docx_path)
    print(f"wrote {docx_path}")


if __name__ == "__main__":
    main()
